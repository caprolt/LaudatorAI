"""Resume processing service with Celery tasks."""

import json
import os
import tempfile
import time
from typing import Dict, Any, List, Optional
from pathlib import Path

import pdfplumber
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from weasyprint import HTML, CSS
from weasyprint.text.fonts import FontConfiguration

from app.core.celery_app import celery_app
from app.core.logging import log_task_start, log_task_complete, log_task_error
from app.services.file_storage import file_storage
from app.templates.default_resume_template import get_template


class ResumeParser:
    """Parse resume files into structured JSON."""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.doc']
    
    def parse_resume(self, file_path: str) -> Dict[str, Any]:
        """Parse a resume file and extract structured content."""
        file_ext = Path(file_path).suffix.lower()
        
        if file_ext not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_ext}")
        
        if file_ext == '.pdf':
            return self._parse_pdf(file_path)
        elif file_ext in ['.docx', '.doc']:
            return self._parse_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")
    
    def _parse_pdf(self, file_path: str) -> Dict[str, Any]:
        """Parse PDF resume."""
        text_content = ""
        
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text_content += page.extract_text() or ""
        
        return self._extract_structured_content(text_content)
    
    def _parse_docx(self, file_path: str) -> Dict[str, Any]:
        """Parse DOCX resume."""
        doc = Document(file_path)
        text_content = ""
        
        for paragraph in doc.paragraphs:
            text_content += paragraph.text + "\n"
        
        return self._extract_structured_content(text_content)
    
    def _extract_structured_content(self, text: str) -> Dict[str, Any]:
        """Extract structured content from text using basic parsing rules."""
        lines = text.split('\n')
        structured_content = {
            "personal_info": {},
            "summary": "",
            "experience": [],
            "education": [],
            "skills": [],
            "certifications": [],
            "projects": [],
            "languages": [],
            "raw_text": text
        }
        
        current_section = None
        current_item = {}
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Detect sections
            lower_line = line.lower()
            if any(keyword in lower_line for keyword in ['experience', 'work history', 'employment']):
                current_section = 'experience'
                continue
            elif any(keyword in lower_line for keyword in ['education', 'academic', 'degree']):
                current_section = 'education'
                continue
            elif any(keyword in lower_line for keyword in ['skills', 'technical skills', 'competencies']):
                current_section = 'skills'
                continue
            elif any(keyword in lower_line for keyword in ['certifications', 'certificates']):
                current_section = 'certifications'
                continue
            elif any(keyword in lower_line for keyword in ['projects', 'portfolio']):
                current_section = 'projects'
                continue
            elif any(keyword in lower_line for keyword in ['languages']):
                current_section = 'languages'
                continue
            elif any(keyword in lower_line for keyword in ['summary', 'objective', 'profile']):
                current_section = 'summary'
                continue
            
            # Extract personal info (email, phone, location)
            if '@' in line and '.' in line:
                structured_content['personal_info']['email'] = line
            elif any(char.isdigit() for char in line) and len(line.replace(' ', '').replace('-', '').replace('(', '').replace(')', '')) >= 10:
                structured_content['personal_info']['phone'] = line
            elif any(keyword in lower_line for keyword in ['street', 'avenue', 'road', 'drive', 'lane']):
                structured_content['personal_info']['address'] = line
            
            # Process sections
            if current_section == 'summary':
                structured_content['summary'] += line + " "
            elif current_section == 'skills':
                skills = [skill.strip() for skill in line.split(',')]
                structured_content['skills'].extend(skills)
            elif current_section == 'experience':
                # Basic experience parsing
                if ' - ' in line or ' at ' in line:
                    if current_item:
                        structured_content['experience'].append(current_item)
                    current_item = {'title': line}
                elif current_item:
                    if 'description' not in current_item:
                        current_item['description'] = line
                    else:
                        current_item['description'] += " " + line
            elif current_section == 'education':
                if current_item:
                    structured_content['education'].append(current_item)
                current_item = {'institution': line}
        
        # Add final items
        if current_section == 'experience' and current_item:
            structured_content['experience'].append(current_item)
        elif current_section == 'education' and current_item:
            structured_content['education'].append(current_item)
        
        return structured_content


class ResumeTailor:
    """Tailor resume content for specific job descriptions."""
    
    def __init__(self):
        self.parser = ResumeParser()
    
    def tailor_resume(self, resume_content: Dict[str, Any], job_description: Dict[str, Any]) -> Dict[str, Any]:
        """Tailor resume content based on job description."""
        tailored_content = resume_content.copy()
        
        # Extract key requirements from job description
        job_requirements = self._extract_job_requirements(job_description)
        
        # Tailor skills section
        tailored_content['skills'] = self._tailor_skills(
            resume_content.get('skills', []), 
            job_requirements.get('skills', [])
        )
        
        # Tailor experience section
        tailored_content['experience'] = self._tailor_experience(
            resume_content.get('experience', []),
            job_requirements
        )
        
        # Tailor summary
        tailored_content['summary'] = self._tailor_summary(
            resume_content.get('summary', ''),
            job_requirements
        )
        
        return tailored_content
    
    def _extract_job_requirements(self, job_description: Dict[str, Any]) -> Dict[str, Any]:
        """Extract key requirements from job description."""
        requirements = {
            'skills': [],
            'keywords': [],
            'responsibilities': []
        }
        
        # Extract from normalized content if available
        normalized = job_description.get('normalized_content', {})
        if isinstance(normalized, str):
            try:
                normalized = json.loads(normalized)
            except:
                normalized = {}
        
        # Extract skills and requirements
        if isinstance(normalized, dict):
            requirements['skills'] = normalized.get('required_skills', [])
            requirements['keywords'] = normalized.get('keywords', [])
            requirements['responsibilities'] = normalized.get('responsibilities', [])
        
        return requirements
    
    def _tailor_skills(self, resume_skills: List[str], job_skills: List[str]) -> List[str]:
        """Tailor skills to match job requirements."""
        if not job_skills:
            return resume_skills
        
        # Prioritize skills that match job requirements
        matched_skills = []
        other_skills = []
        
        for skill in resume_skills:
            skill_lower = skill.lower()
            if any(job_skill.lower() in skill_lower or skill_lower in job_skill.lower() 
                   for job_skill in job_skills):
                matched_skills.append(skill)
            else:
                other_skills.append(skill)
        
        # Return matched skills first, then others
        return matched_skills + other_skills
    
    def _tailor_experience(self, experience: List[Dict[str, Any]], job_requirements: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Tailor experience descriptions to match job requirements."""
        tailored_experience = []
        
        for exp in experience:
            tailored_exp = exp.copy()
            description = exp.get('description', '')
            
            # Add relevant keywords to description if not present
            keywords = job_requirements.get('keywords', [])
            for keyword in keywords:
                if keyword.lower() not in description.lower():
                    description += f" • {keyword}"
            
            tailored_exp['description'] = description
            tailored_experience.append(tailored_exp)
        
        return tailored_experience
    
    def _tailor_summary(self, summary: str, job_requirements: Dict[str, Any]) -> str:
        """Tailor summary to match job requirements."""
        if not summary:
            return summary
        
        # Add relevant keywords to summary
        keywords = job_requirements.get('keywords', [])
        for keyword in keywords[:3]:  # Limit to first 3 keywords
            if keyword.lower() not in summary.lower():
                summary += f" Proficient in {keyword}."
        
        return summary


class ResumeGenerator:
    """Generate DOCX and PDF resumes."""
    
    def __init__(self):
        self.templates_dir = Path(__file__).parent.parent / "templates"
        self.templates_dir.mkdir(exist_ok=True)
    
    def generate_docx(self, resume_content: Dict[str, Any], template_name: str = "default") -> str:
        """Generate DOCX resume from structured content."""
        doc = Document()
        
        # Get template configuration
        template = get_template(template_name)
        
        # Set up document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = template.margins['top']
            section.bottom_margin = template.margins['bottom']
            section.left_margin = template.margins['left']
            section.right_margin = template.margins['right']
        
        # Add sections based on template order
        for section_name in template.sections:
            if section_name == 'header':
                self._add_header(doc, resume_content.get('personal_info', {}), template)
            elif section_name == 'summary' and resume_content.get('summary'):
                self._add_summary(doc, resume_content['summary'], template)
            elif section_name == 'experience' and resume_content.get('experience'):
                self._add_experience(doc, resume_content['experience'], template)
            elif section_name == 'education' and resume_content.get('education'):
                self._add_education(doc, resume_content['education'], template)
            elif section_name == 'skills' and resume_content.get('skills'):
                self._add_skills(doc, resume_content['skills'], template)
            elif section_name == 'certifications' and resume_content.get('certifications'):
                self._add_certifications(doc, resume_content['certifications'], template)
            elif section_name == 'projects' and resume_content.get('projects'):
                self._add_projects(doc, resume_content['projects'], template)
        
        # Save to temporary file
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_file.name)
        return temp_file.name
    
    def _add_header(self, doc: Document, personal_info: Dict[str, Any], template):
        """Add header with personal information."""
        config = template.get_section_config('header')
        formatted_info = template.format_personal_info(personal_info)
        
        for i, line in enumerate(formatted_info):
            para = doc.add_paragraph()
            run = para.add_run(line)
            
            # Apply font settings
            if i == 0:  # Name
                font_config = template.fonts['name']
                run.font.name = font_config['name']
                run.font.size = font_config['size']
                run.bold = font_config['bold']
            else:  # Contact info
                font_config = template.fonts['contact']
                run.font.name = font_config['name']
                run.font.size = font_config['size']
            
            # Apply alignment
            para.alignment = config['alignment']
            
            # Apply spacing
            if 'spacing_after' in config:
                para.paragraph_format.space_after = config['spacing_after']
    
    def _add_summary(self, doc: Document, summary: str, template):
        """Add summary section."""
        config = template.get_section_config('summary')
        
        heading = doc.add_heading("Summary", level=1)
        heading.alignment = config.get('alignment', WD_ALIGN_PARAGRAPH.LEFT)
        
        para = doc.add_paragraph(summary)
        para.alignment = config.get('alignment', WD_ALIGN_PARAGRAPH.LEFT)
        
        if 'spacing_after' in config:
            para.paragraph_format.space_after = config['spacing_after']
    
    def _add_experience(self, doc: Document, experience: List[Dict[str, Any]], template):
        """Add experience section."""
        config = template.get_section_config('experience')
        
        heading = doc.add_heading("Experience", level=1)
        heading.alignment = config.get('alignment', WD_ALIGN_PARAGRAPH.LEFT)
        
        for exp in experience:
            # Job title and company
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(exp.get('title', 'Unknown Position'))
            title_run.bold = True
            title_para.alignment = config.get('alignment', WD_ALIGN_PARAGRAPH.LEFT)
            
            # Description
            if exp.get('description'):
                desc_para = doc.add_paragraph(exp['description'])
                desc_para.alignment = config.get('alignment', WD_ALIGN_PARAGRAPH.LEFT)
            
            # Add spacing between items
            if 'item_spacing' in config:
                doc.add_paragraph().paragraph_format.space_after = config['item_spacing']
        
        if 'spacing_after' in config:
            doc.paragraphs[-1].paragraph_format.space_after = config['spacing_after']
    
    def _add_education(self, doc: Document, education: List[Dict[str, Any]], template):
        """Add education section."""
        config = template.get_section_config('education')
        
        heading = doc.add_heading("Education", level=1)
        heading.alignment = config.get('alignment', WD_ALIGN_PARAGRAPH.LEFT)
        
        for edu in education:
            institution_para = doc.add_paragraph()
            institution_run = institution_para.add_run(edu.get('institution', 'Unknown Institution'))
            institution_run.bold = True
            institution_para.alignment = config.get('alignment', WD_ALIGN_PARAGRAPH.LEFT)
            
            doc.add_paragraph()  # Spacing
        
        if 'spacing_after' in config:
            doc.paragraphs[-1].paragraph_format.space_after = config['spacing_after']
    
    def _add_skills(self, doc: Document, skills: List[str], template):
        """Add skills section."""
        config = template.get_section_config('skills')
        
        heading = doc.add_heading("Skills", level=1)
        heading.alignment = config.get('alignment', WD_ALIGN_PARAGRAPH.LEFT)
        
        # Format skills using template
        skills_text = template.format_skills(skills)
        para = doc.add_paragraph(skills_text)
        para.alignment = config.get('alignment', WD_ALIGN_PARAGRAPH.LEFT)
        
        if 'spacing_after' in config:
            para.paragraph_format.space_after = config['spacing_after']
    
    def _add_certifications(self, doc: Document, certifications: List[str], template):
        """Add certifications section."""
        config = template.get_section_config('certifications')
        
        heading = doc.add_heading("Certifications", level=1)
        heading.alignment = config.get('alignment', WD_ALIGN_PARAGRAPH.LEFT)
        
        for cert in certifications:
            para = doc.add_paragraph(cert)
            para.alignment = config.get('alignment', WD_ALIGN_PARAGRAPH.LEFT)
        
        if 'spacing_after' in config:
            doc.paragraphs[-1].paragraph_format.space_after = config['spacing_after']
    
    def _add_projects(self, doc: Document, projects: List[Dict[str, Any]], template):
        """Add projects section."""
        config = template.get_section_config('projects')
        
        heading = doc.add_heading("Projects", level=1)
        heading.alignment = config.get('alignment', WD_ALIGN_PARAGRAPH.LEFT)
        
        for project in projects:
            # Project title
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(project.get('title', 'Unknown Project'))
            title_run.bold = True
            title_para.alignment = config.get('alignment', WD_ALIGN_PARAGRAPH.LEFT)
            
            # Project description
            if project.get('description'):
                desc_para = doc.add_paragraph(project['description'])
                desc_para.alignment = config.get('alignment', WD_ALIGN_PARAGRAPH.LEFT)
            
            doc.add_paragraph()  # Spacing
        
        if 'spacing_after' in config:
            doc.paragraphs[-1].paragraph_format.space_after = config['spacing_after']
    
    def generate_pdf(self, docx_path: str) -> str:
        """Convert DOCX to PDF using weasyprint."""
        # First convert DOCX to HTML (simplified approach)
        html_content = self._docx_to_html(docx_path)
        
        # Generate PDF from HTML
        font_config = FontConfiguration()
        html_doc = HTML(string=html_content)
        css = CSS(string='''
            body { font-family: Arial, sans-serif; margin: 1in; }
            h1 { color: #2c3e50; border-bottom: 2px solid #3498db; }
            .header { text-align: center; margin-bottom: 20px; }
            .name { font-size: 24px; font-weight: bold; }
            .contact { color: #7f8c8d; }
        ''', font_config=font_config)
        
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
        html_doc.write_pdf(temp_file.name, stylesheets=[css], font_config=font_config)
        return temp_file.name
    
    def _docx_to_html(self, docx_path: str) -> str:
        """Convert DOCX to HTML (simplified)."""
        doc = Document(docx_path)
        html_parts = ['<html><body>']
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                if paragraph.style.name.startswith('Heading'):
                    html_parts.append(f'<h1>{paragraph.text}</h1>')
                else:
                    html_parts.append(f'<p>{paragraph.text}</p>')
        
        html_parts.append('</body></html>')
        return '\n'.join(html_parts)


# Global instances
resume_parser = ResumeParser()
resume_tailor = ResumeTailor()
resume_generator = ResumeGenerator()


@celery_app.task(bind=True)
def parse_resume(self, resume_id: int, file_path: str) -> Dict[str, Any]:
    """Parse a resume file and extract structured content."""
    task_id = self.request.id
    task_type = "resume_parsing"
    
    try:
        log_task_start(task_id, task_type, resume_id=resume_id, file_path=file_path)
        start_time = time.time()
        
        # Download file from storage
        local_file_path = file_storage.download_file(file_path)
        
        try:
            # Parse resume
            parsed_content = resume_parser.parse_resume(local_file_path)
            
            # Update database
            from app.models import Resume
            from app.core.database import SessionLocal
            
            db = SessionLocal()
            try:
                resume = db.query(Resume).filter(Resume.id == resume_id).first()
                if resume:
                    resume.parsed_content = json.dumps(parsed_content)
                    resume.status = "parsed"
                    db.commit()
            finally:
                db.close()
            
            result = {
                "resume_id": resume_id,
                "status": "parsed",
                "parsed_content": parsed_content,
                "message": "Resume parsed successfully"
            }
            
        finally:
            # Clean up local file
            if os.path.exists(local_file_path):
                os.unlink(local_file_path)
        
        duration = time.time() - start_time
        log_task_complete(task_id, task_type, duration=duration, resume_id=resume_id)
        
        return result
        
    except Exception as e:
        duration = time.time() - start_time
        log_task_error(task_id, task_type, str(e), resume_id=resume_id, duration=duration)
        raise


@celery_app.task(bind=True)
def tailor_resume(self, application_id: int, job_id: int, resume_id: int) -> Dict[str, Any]:
    """Tailor a resume for a specific job posting."""
    task_id = self.request.id
    task_type = "resume_tailoring"
    
    try:
        log_task_start(task_id, task_type, application_id=application_id, job_id=job_id, resume_id=resume_id)
        start_time = time.time()
        
        # Get data from database
        from app.models import JobApplication, Job, Resume
        from app.core.database import SessionLocal
        
        db = SessionLocal()
        try:
            application = db.query(JobApplication).filter(JobApplication.id == application_id).first()
            job = db.query(Job).filter(Job.id == job_id).first()
            resume = db.query(Resume).filter(Resume.id == resume_id).first()
            
            if not all([application, job, resume]):
                raise ValueError("Application, job, or resume not found")
            
            # Parse resume content
            resume_content = json.loads(resume.parsed_content) if resume.parsed_content else {}
            
            # Parse job description
            job_description = {
                'normalized_content': job.normalized_content,
                'description': job.description,
                'requirements': job.requirements
            }
            
            # Tailor resume
            tailored_content = resume_tailor.tailor_resume(resume_content, job_description)
            
            # Generate DOCX
            docx_path = resume_generator.generate_docx(tailored_content)
            
            try:
                # Upload tailored resume
                object_name = f"applications/{application_id}/tailored_resume.docx"
                tailored_resume_path = file_storage.upload_file(docx_path, object_name)
                
                # Generate PDF
                pdf_path = resume_generator.generate_pdf(docx_path)
                
                try:
                    # Upload PDF
                    pdf_object_name = f"applications/{application_id}/tailored_resume.pdf"
                    tailored_resume_pdf_path = file_storage.upload_file(pdf_path, pdf_object_name)
                    
                    # Update application
                    application.tailored_resume_path = tailored_resume_path
                    application.status = "completed"
                    db.commit()
                    
                    result = {
                        "application_id": application_id,
                        "job_id": job_id,
                        "resume_id": resume_id,
                        "tailored_resume_path": tailored_resume_path,
                        "tailored_resume_pdf_path": tailored_resume_pdf_path,
                        "status": "tailored",
                        "message": "Resume tailored successfully"
                    }
                    
                finally:
                    if os.path.exists(pdf_path):
                        os.unlink(pdf_path)
                        
            finally:
                if os.path.exists(docx_path):
                    os.unlink(docx_path)
                    
        finally:
            db.close()
        
        duration = time.time() - start_time
        log_task_complete(task_id, task_type, duration=duration, application_id=application_id)
        
        return result
        
    except Exception as e:
        duration = time.time() - start_time
        log_task_error(task_id, task_type, str(e), application_id=application_id, duration=duration)
        raise


@celery_app.task(bind=True)
def generate_resume_preview(self, resume_id: int, job_id: Optional[int] = None) -> Dict[str, Any]:
    """Generate a preview of the resume (HTML format for web display)."""
    task_id = self.request.id
    task_type = "resume_preview"
    
    try:
        log_task_start(task_id, task_type, resume_id=resume_id, job_id=job_id)
        start_time = time.time()
        
        # Get data from database
        from app.models import Resume, Job
        from app.core.database import SessionLocal
        
        db = SessionLocal()
        try:
            resume = db.query(Resume).filter(Resume.id == resume_id).first()
            job = db.query(Job).filter(Job.id == job_id).first() if job_id else None
            
            if not resume:
                raise ValueError("Resume not found")
            
            # Parse resume content
            resume_content = json.loads(resume.parsed_content) if resume.parsed_content else {}
            
            # Tailor if job is provided
            if job:
                job_description = {
                    'normalized_content': job.normalized_content,
                    'description': job.description,
                    'requirements': job.requirements
                }
                resume_content = resume_tailor.tailor_resume(resume_content, job_description)
            
            # Generate HTML preview
            html_preview = _generate_html_preview(resume_content)
            
            result = {
                "resume_id": resume_id,
                "job_id": job_id,
                "html_preview": html_preview,
                "status": "preview_generated",
                "message": "Resume preview generated successfully"
            }
            
        finally:
            db.close()
        
        duration = time.time() - start_time
        log_task_complete(task_id, task_type, duration=duration, resume_id=resume_id)
        
        return result
        
    except Exception as e:
        duration = time.time() - start_time
        log_task_error(task_id, task_type, str(e), resume_id=resume_id, duration=duration)
        raise


def _generate_html_preview(resume_content: Dict[str, Any]) -> str:
    """Generate HTML preview of resume content."""
    html_parts = ['<div class="resume-preview">']
    
    # Personal info
    personal_info = resume_content.get('personal_info', {})
    if personal_info:
        html_parts.append('<div class="header">')
        if personal_info.get('name'):
            html_parts.append(f'<h1>{personal_info["name"]}</h1>')
        contact_info = []
        if personal_info.get('email'):
            contact_info.append(personal_info['email'])
        if personal_info.get('phone'):
            contact_info.append(personal_info['phone'])
        if contact_info:
            html_parts.append(f'<p class="contact">{", ".join(contact_info)}</p>')
        html_parts.append('</div>')
    
    # Summary
    if resume_content.get('summary'):
        html_parts.append('<section class="summary">')
        html_parts.append('<h2>Summary</h2>')
        html_parts.append(f'<p>{resume_content["summary"]}</p>')
        html_parts.append('</section>')
    
    # Experience
    if resume_content.get('experience'):
        html_parts.append('<section class="experience">')
        html_parts.append('<h2>Experience</h2>')
        for exp in resume_content['experience']:
            html_parts.append('<div class="experience-item">')
            if exp.get('title'):
                html_parts.append(f'<h3>{exp["title"]}</h3>')
            if exp.get('description'):
                html_parts.append(f'<p>{exp["description"]}</p>')
            html_parts.append('</div>')
        html_parts.append('</section>')
    
    # Skills
    if resume_content.get('skills'):
        html_parts.append('<section class="skills">')
        html_parts.append('<h2>Skills</h2>')
        html_parts.append(f'<p>{", ".join(resume_content["skills"])}</p>')
        html_parts.append('</section>')
    
    html_parts.append('</div>')
    return '\n'.join(html_parts)
