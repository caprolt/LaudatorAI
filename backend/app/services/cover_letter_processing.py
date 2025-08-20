"""Cover letter processing service with Celery tasks."""

import json
import os
import tempfile
import time
from typing import Dict, Any, List, Optional
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
# Optional WeasyPrint import for PDF generation
try:
    from weasyprint import HTML, CSS
    from weasyprint.text.fonts import FontConfiguration
    WEASYPRINT_AVAILABLE = True
except (ImportError, OSError):
    WEASYPRINT_AVAILABLE = False
    HTML = None
    CSS = None
    FontConfiguration = None
import openai
from openai import OpenAI

from app.core.celery_app import celery_app
from app.core.config import settings
from app.core.logging import log_task_start, log_task_complete, log_task_error
from app.services.file_storage import file_storage
from app.templates.default_cover_letter_template import get_template


class CoverLetterGenerator:
    """Generate tailored cover letters using LLM."""
    
    def __init__(self):
        self.client = None
        self._setup_llm_client()
    
    def _setup_llm_client(self):
        """Setup LLM client based on configuration."""
        if settings.LLM_PROVIDER == "openai":
            if not settings.OPENAI_API_KEY:
                raise ValueError("OpenAI API key not configured")
            self.client = OpenAI(api_key=settings.OPENAI_API_KEY)
        else:
            # TODO: Add support for Ollama and HuggingFace
            raise NotImplementedError(f"LLM provider {settings.LLM_PROVIDER} not implemented")
    
    def generate_cover_letter_content(
        self, 
        job_description: Dict[str, Any], 
        resume_data: Dict[str, Any],
        personal_info: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Generate cover letter content using LLM."""
        
        # Prepare the prompt
        prompt = self._build_cover_letter_prompt(job_description, resume_data, personal_info)
        
        try:
            response = self.client.chat.completions.create(
                model=settings.LLM_MODEL,  # Use configured model
                messages=[
                    {
                        "role": "system",
                        "content": "You are an expert cover letter writer. Create compelling, professional cover letters that highlight relevant experience and skills for specific job opportunities."
                    },
                    {
                        "role": "user",
                        "content": prompt
                    }
                ],
                temperature=0.7,
                max_tokens=1500
            )
            
            content = response.choices[0].message.content
            
            # Parse the response into structured format
            return self._parse_cover_letter_content(content, personal_info)
            
        except Exception as e:
            raise Exception(f"Failed to generate cover letter content: {str(e)}")
    
    def _build_cover_letter_prompt(
        self, 
        job_description: Dict[str, Any], 
        resume_data: Dict[str, Any],
        personal_info: Dict[str, Any]
    ) -> str:
        """Build the prompt for cover letter generation."""
        
        # Extract key information
        job_title = job_description.get('title', '')
        company_name = job_description.get('company', '')
        job_summary = job_description.get('summary', '')
        key_requirements = job_description.get('requirements', [])
        
        # Extract relevant experience and skills
        experience = resume_data.get('experience', [])
        skills = resume_data.get('skills', [])
        education = resume_data.get('education', [])
        
        prompt = f"""
        Please write a professional cover letter for the following job opportunity:
        
        Job Title: {job_title}
        Company: {company_name}
        Job Summary: {job_summary}
        
        Key Requirements:
        {chr(10).join([f"- {req}" for req in key_requirements[:10]])}
        
        Candidate Information:
        Name: {personal_info.get('name', '')}
        Email: {personal_info.get('email', '')}
        Phone: {personal_info.get('phone', '')}
        
        Relevant Experience:
        {chr(10).join([f"- {exp.get('title', '')} at {exp.get('company', '')} ({exp.get('duration', '')})" for exp in experience[:3]])}
        
        Key Skills:
        {chr(10).join([f"- {skill}" for skill in skills[:10]])}
        
        Education:
        {chr(10).join([f"- {edu.get('degree', '')} from {edu.get('institution', '')}" for edu in education[:2]])}
        
        Please create a cover letter that:
        1. Addresses the hiring manager professionally
        2. Opens with a compelling introduction that mentions the specific position
        3. Highlights 2-3 most relevant experiences that match the job requirements
        4. Demonstrates understanding of the company and role
        5. Closes with enthusiasm and a call to action
        6. Is concise (3-4 paragraphs, approximately 300-400 words)
        7. Maintains a professional yet engaging tone
        
        Format the response as JSON with the following structure:
        {{
            "greeting": "Dear [Hiring Manager/Recruiter]",
            "opening": "Opening paragraph...",
            "body": "Body paragraphs...",
            "closing": "Closing paragraph...",
            "signature": "Sincerely,\\n[Name]"
        }}
        """
        
        return prompt
    
    def _parse_cover_letter_content(self, content: str, personal_info: Dict[str, Any]) -> Dict[str, Any]:
        """Parse the LLM response into structured cover letter content."""
        try:
            # Try to parse as JSON first
            parsed_content = json.loads(content)
            
            # Ensure all required fields are present
            required_fields = ['greeting', 'opening', 'body', 'closing', 'signature']
            for field in required_fields:
                if field not in parsed_content:
                    parsed_content[field] = ""
            
            # Replace placeholder with actual name
            if '[Name]' in parsed_content['signature']:
                parsed_content['signature'] = parsed_content['signature'].replace('[Name]', personal_info.get('name', ''))
            
            return parsed_content
            
        except json.JSONDecodeError:
            # Fallback: parse as plain text
            paragraphs = content.split('\n\n')
            return {
                "greeting": "Dear Hiring Manager,",
                "opening": paragraphs[0] if len(paragraphs) > 0 else "",
                "body": "\n\n".join(paragraphs[1:-1]) if len(paragraphs) > 2 else "",
                "closing": paragraphs[-1] if len(paragraphs) > 1 else "",
                "signature": f"Sincerely,\n{personal_info.get('name', '')}"
            }


class CoverLetterDocumentGenerator:
    """Generate DOCX and PDF cover letter documents."""
    
    def __init__(self):
        self.template = get_template()
    
    def generate_docx(
        self, 
        cover_letter_content: Dict[str, Any], 
        job_description: Dict[str, Any],
        personal_info: Dict[str, Any]
    ) -> bytes:
        """Generate DOCX cover letter."""
        doc = Document()
        
        # Set up document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Add header with personal info
        self._add_header(doc, personal_info)
        
        # Add date
        self._add_date(doc)
        
        # Add recipient info
        self._add_recipient_info(doc, job_description)
        
        # Add greeting
        self._add_greeting(doc, cover_letter_content.get('greeting', ''))
        
        # Add content
        self._add_content(doc, cover_letter_content)
        
        # Add signature
        self._add_signature(doc, cover_letter_content.get('signature', ''))
        
        # Save to bytes
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_file:
            doc.save(tmp_file.name)
            with open(tmp_file.name, 'rb') as f:
                content = f.read()
            os.unlink(tmp_file.name)
        
        return content
    
    def generate_pdf(
        self, 
        cover_letter_content: Dict[str, Any], 
        job_description: Dict[str, Any],
        personal_info: Dict[str, Any]
    ) -> bytes:
        """Generate PDF cover letter."""
        if not WEASYPRINT_AVAILABLE:
            raise RuntimeError("WeasyPrint is not available. PDF generation requires WeasyPrint to be installed with proper system dependencies.")
        
        # Generate HTML content
        html_content = self._generate_html(cover_letter_content, job_description, personal_info)
        
        # Convert to PDF
        font_config = FontConfiguration()
        css = CSS(string=self.template['css'], font_config=font_config)
        
        html_doc = HTML(string=html_content)
        pdf_bytes = html_doc.write_pdf(stylesheets=[css], font_config=font_config)
        
        return pdf_bytes
    
    def _add_header(self, doc: Document, personal_info: Dict[str, Any]):
        """Add header with personal information."""
        header_para = doc.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        name_run = header_para.add_run(personal_info.get('name', ''))
        name_run.bold = True
        name_run.font.size = Pt(16)
        
        # Add contact info
        contact_info = []
        if personal_info.get('email'):
            contact_info.append(personal_info['email'])
        if personal_info.get('phone'):
            contact_info.append(personal_info['phone'])
        if personal_info.get('location'):
            contact_info.append(personal_info['location'])
        
        if contact_info:
            contact_para = doc.add_paragraph()
            contact_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            contact_run = contact_para.add_run('\n'.join(contact_info))
            contact_run.font.size = Pt(10)
        
        doc.add_paragraph()  # Add spacing
    
    def _add_date(self, doc: Document):
        """Add current date."""
        from datetime import datetime
        
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        date_run = date_para.add_run(datetime.now().strftime("%B %d, %Y"))
        date_run.font.size = Pt(10)
        
        doc.add_paragraph()  # Add spacing
    
    def _add_recipient_info(self, doc: Document, job_description: Dict[str, Any]):
        """Add recipient information."""
        company_name = job_description.get('company', '')
        
        recipient_para = doc.add_paragraph()
        recipient_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        recipient_run = recipient_para.add_run(f"Hiring Manager\n{company_name}")
        recipient_run.font.size = Pt(10)
        
        doc.add_paragraph()  # Add spacing
    
    def _add_greeting(self, doc: Document, greeting: str):
        """Add greeting."""
        greeting_para = doc.add_paragraph()
        greeting_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        greeting_run = greeting_para.add_run(greeting)
        greeting_run.font.size = Pt(10)
        
        doc.add_paragraph()  # Add spacing
    
    def _add_content(self, doc: Document, cover_letter_content: Dict[str, Any]):
        """Add main content."""
        # Add opening paragraph
        if cover_letter_content.get('opening'):
            opening_para = doc.add_paragraph()
            opening_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            opening_run = opening_para.add_run(cover_letter_content['opening'])
            opening_run.font.size = Pt(10)
            doc.add_paragraph()  # Add spacing
        
        # Add body paragraphs
        if cover_letter_content.get('body'):
            body_paragraphs = cover_letter_content['body'].split('\n\n')
            for paragraph_text in body_paragraphs:
                if paragraph_text.strip():
                    body_para = doc.add_paragraph()
                    body_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    body_run = body_para.add_run(paragraph_text.strip())
                    body_run.font.size = Pt(10)
                    doc.add_paragraph()  # Add spacing
    
    def _add_signature(self, doc: Document, signature: str):
        """Add signature."""
        signature_para = doc.add_paragraph()
        signature_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        signature_run = signature_para.add_run(signature)
        signature_run.font.size = Pt(10)
    
    def _generate_html(
        self, 
        cover_letter_content: Dict[str, Any], 
        job_description: Dict[str, Any],
        personal_info: Dict[str, Any]
    ) -> str:
        """Generate HTML content for PDF conversion."""
        from datetime import datetime
        
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <title>Cover Letter</title>
        </head>
        <body>
            <div class="cover-letter">
                <div class="header">
                    <div class="name">{personal_info.get('name', '')}</div>
                    <div class="contact-info">
                        {personal_info.get('email', '')}<br>
                        {personal_info.get('phone', '')}<br>
                        {personal_info.get('location', '')}
                    </div>
                </div>
                
                <div class="date">{datetime.now().strftime("%B %d, %Y")}</div>
                
                <div class="recipient">
                    Hiring Manager<br>
                    {job_description.get('company', '')}
                </div>
                
                <div class="greeting">{cover_letter_content.get('greeting', '')}</div>
                
                <div class="content">
                    <p>{cover_letter_content.get('opening', '')}</p>
                    <p>{cover_letter_content.get('body', '')}</p>
                </div>
                
                <div class="signature">{cover_letter_content.get('signature', '')}</div>
            </div>
        </body>
        </html>
        """
        
        return html_content


# Celery tasks
@celery_app.task(bind=True)
def generate_cover_letter(self, application_id: int, job_id: int, resume_id: int) -> Dict[str, Any]:
    """Generate a cover letter for a job application."""
    task_id = self.request.id
    task_type = "cover_letter_generation"
    
    try:
        log_task_start(task_id, task_type, application_id=application_id, job_id=job_id, resume_id=resume_id)
        start_time = time.time()
        
        # TODO: Get job description, resume data, and personal info from database
        # For now, using placeholder data
        job_description = {
            "title": "Software Engineer",
            "company": "Example Corp",
            "summary": "We are looking for a talented software engineer...",
            "requirements": ["Python", "FastAPI", "PostgreSQL"]
        }
        
        resume_data = {
            "experience": [
                {"title": "Software Engineer", "company": "Previous Corp", "duration": "2 years"}
            ],
            "skills": ["Python", "FastAPI", "PostgreSQL"],
            "education": [{"degree": "BS Computer Science", "institution": "University"}]
        }
        
        personal_info = {
            "name": "John Doe",
            "email": "john.doe@example.com",
            "phone": "(555) 123-4567",
            "location": "San Francisco, CA"
        }
        
        # Generate cover letter content
        generator = CoverLetterGenerator()
        cover_letter_content = generator.generate_cover_letter_content(
            job_description, resume_data, personal_info
        )
        
        # Generate documents
        doc_generator = CoverLetterDocumentGenerator()
        docx_content = doc_generator.generate_docx(cover_letter_content, job_description, personal_info)
        pdf_content = doc_generator.generate_pdf(cover_letter_content, job_description, personal_info)
        
        # Store files
        docx_filename = f"cover_letter_{application_id}.docx"
        pdf_filename = f"cover_letter_{application_id}.pdf"
        
        docx_url = file_storage.upload_file(docx_content, docx_filename, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        pdf_url = file_storage.upload_file(pdf_content, pdf_filename, "application/pdf")
        
        result = {
            "application_id": application_id,
            "job_id": job_id,
            "resume_id": resume_id,
            "status": "generated",
            "cover_letter_content": cover_letter_content,
            "docx_url": docx_url,
            "pdf_url": pdf_url,
            "message": "Cover letter generated successfully"
        }
        
        duration = time.time() - start_time
        log_task_complete(task_id, task_type, duration=duration, application_id=application_id)
        
        return result
        
    except Exception as e:
        duration = time.time() - start_time
        log_task_error(task_id, task_type, str(e), application_id=application_id, duration=duration)
        raise


@celery_app.task(bind=True)
def preview_cover_letter(self, job_description: Dict[str, Any], resume_data: Dict[str, Any], personal_info: Dict[str, Any]) -> Dict[str, Any]:
    """Generate a preview cover letter without saving to storage."""
    task_id = self.request.id
    task_type = "cover_letter_preview"
    
    try:
        log_task_start(task_id, task_type)
        start_time = time.time()
        
        # Generate cover letter content
        generator = CoverLetterGenerator()
        cover_letter_content = generator.generate_cover_letter_content(
            job_description, resume_data, personal_info
        )
        
        result = {
            "cover_letter_content": cover_letter_content,
            "status": "preview_generated",
            "message": "Cover letter preview generated successfully"
        }
        
        duration = time.time() - start_time
        log_task_complete(task_id, task_type, duration=duration)
        
        return result
        
    except Exception as e:
        duration = time.time() - start_time
        log_task_error(task_id, task_type, str(e), duration=duration)
        raise
