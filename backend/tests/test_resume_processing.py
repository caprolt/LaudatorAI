"""Tests for resume processing functionality."""

import json
import tempfile
import os
from unittest.mock import Mock, patch

import pytest
from docx import Document

from app.services.resume_processing import ResumeParser, ResumeTailor, ResumeGenerator
from app.templates.default_resume_template import get_template


class TestResumeParser:
    """Test resume parsing functionality."""
    
    def test_parse_pdf(self):
        """Test PDF parsing."""
        parser = ResumeParser()
        
        # Create a mock PDF file
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_file:
            temp_file.write(b"Mock PDF content")
            temp_file_path = temp_file.name
        
        try:
            with patch('pdfplumber.open') as mock_pdf:
                mock_page = Mock()
                mock_page.extract_text.return_value = "Software Engineer at Tech Corp\nPython, JavaScript\nBachelor's in Computer Science"
                mock_pdf.return_value.__enter__.return_value.pages = [mock_page]
                
                result = parser.parse_resume(temp_file_path)
                
                assert 'personal_info' in result
                assert 'experience' in result
                assert 'skills' in result
                assert 'education' in result
        finally:
            os.unlink(temp_file_path)
    
    def test_parse_docx(self):
        """Test DOCX parsing."""
        parser = ResumeParser()
        
        # Create a mock DOCX file
        doc = Document()
        doc.add_paragraph("Software Engineer at Tech Corp")
        doc.add_paragraph("Python, JavaScript")
        doc.add_paragraph("Bachelor's in Computer Science")
        
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
            doc.save(temp_file.name)
            temp_file_path = temp_file.name
        
        try:
            result = parser.parse_resume(temp_file_path)
            
            assert 'personal_info' in result
            assert 'experience' in result
            assert 'skills' in result
            assert 'education' in result
        finally:
            os.unlink(temp_file_path)
    
    def test_unsupported_format(self):
        """Test handling of unsupported file formats."""
        parser = ResumeParser()
        
        with tempfile.NamedTemporaryFile(suffix='.txt', delete=False) as temp_file:
            temp_file_path = temp_file.name
        
        try:
            with pytest.raises(ValueError, match="Unsupported file format"):
                parser.parse_resume(temp_file_path)
        finally:
            os.unlink(temp_file_path)


class TestResumeTailor:
    """Test resume tailoring functionality."""
    
    def test_tailor_skills(self):
        """Test skills tailoring."""
        tailor = ResumeTailor()
        
        resume_content = {
            'skills': ['Python', 'JavaScript', 'SQL', 'React']
        }
        
        job_description = {
            'normalized_content': json.dumps({
                'required_skills': ['Python', 'React'],
                'keywords': ['machine learning', 'frontend']
            })
        }
        
        result = tailor.tailor_resume(resume_content, job_description)
        
        # Python and React should be prioritized
        assert result['skills'][0] in ['Python', 'React']
        assert result['skills'][1] in ['Python', 'React']
    
    def test_tailor_experience(self):
        """Test experience tailoring."""
        tailor = ResumeTailor()
        
        resume_content = {
            'experience': [
                {
                    'title': 'Software Engineer',
                    'description': 'Developed web applications'
                }
            ]
        }
        
        job_description = {
            'normalized_content': json.dumps({
                'keywords': ['Python', 'React']
            })
        }
        
        result = tailor.tailor_resume(resume_content, job_description)
        
        # Keywords should be added to experience descriptions
        assert 'Python' in result['experience'][0]['description'] or 'React' in result['experience'][0]['description']


class TestResumeGenerator:
    """Test resume generation functionality."""
    
    def test_generate_docx(self):
        """Test DOCX generation."""
        generator = ResumeGenerator()
        
        resume_content = {
            'personal_info': {
                'name': 'John Doe',
                'email': 'john@example.com'
            },
            'summary': 'Experienced software engineer',
            'experience': [
                {
                    'title': 'Software Engineer',
                    'description': 'Developed web applications'
                }
            ],
            'skills': ['Python', 'JavaScript', 'React']
        }
        
        docx_path = generator.generate_docx(resume_content)
        
        try:
            # Verify the file was created
            assert os.path.exists(docx_path)
            
            # Verify it's a valid DOCX file
            doc = Document(docx_path)
            assert len(doc.paragraphs) > 0
            
            # Check that content was added
            text_content = '\n'.join([p.text for p in doc.paragraphs])
            assert 'John Doe' in text_content
            assert 'Software Engineer' in text_content
        finally:
            if os.path.exists(docx_path):
                os.unlink(docx_path)
    
    def test_generate_pdf(self):
        """Test PDF generation."""
        generator = ResumeGenerator()
        
        # Create a simple DOCX first
        doc = Document()
        doc.add_paragraph("Test resume content")
        
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
            doc.save(temp_file.name)
            docx_path = temp_file.name
        
        try:
            pdf_path = generator.generate_pdf(docx_path)
            
            # Verify the PDF was created
            assert os.path.exists(pdf_path)
            assert pdf_path.endswith('.pdf')
        finally:
            if os.path.exists(docx_path):
                os.unlink(docx_path)
            if os.path.exists(pdf_path):
                os.unlink(pdf_path)


class TestResumeTemplate:
    """Test resume template functionality."""
    
    def test_get_template(self):
        """Test template retrieval."""
        template = get_template('default')
        assert template.name == 'default'
        assert template.description == 'Professional resume template with clean formatting'
    
    def test_template_sections(self):
        """Test template section configuration."""
        template = get_template('default')
        
        assert 'header' in template.sections
        assert 'summary' in template.sections
        assert 'experience' in template.sections
        assert 'skills' in template.sections
    
    def test_format_skills(self):
        """Test skills formatting."""
        template = get_template('default')
        
        skills = ['Python', 'JavaScript', 'SQL', 'AWS', 'Docker']
        formatted = template.format_skills(skills)
        
        assert 'Programming' in formatted
        assert 'Python' in formatted
        assert 'JavaScript' in formatted
        assert 'Cloud/DevOps' in formatted
        assert 'AWS' in formatted
